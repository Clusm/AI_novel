import os
import json
from docx import Document
from docx.shared import Inches
from ebooklib import epub
from datetime import datetime

from src.project import load_project_config, load_outline, list_generated_chapters, load_chapter


def export_to_txt(project_name, output_path=None):
    """导出为TXT格式"""
    if output_path is None:
        output_path = f"{project_name}_{datetime.now().strftime('%Y%m%d')}.txt"
    
    # 加载项目信息
    config = load_project_config(project_name)
    outline = load_outline(project_name)
    chapters = list_generated_chapters(project_name)
    
    # 生成TXT内容
    content = []
    content.append(f"# {config.get('name', project_name)}")
    content.append("")
    content.append("## 大纲")
    content.append(outline)
    content.append("")
    content.append("## 正文")
    content.append("")
    
    for chapter_file in chapters:
        chapter_content = load_chapter(project_name, chapter_file)
        content.append(f"### {chapter_file.replace('.md', '')}")
        content.append(chapter_content)
        content.append("")
    
    # 写入文件
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(content))
    
    return output_path


def export_to_word(project_name, output_path=None):
    """导出为Word格式"""
    if output_path is None:
        output_path = f"{project_name}_{datetime.now().strftime('%Y%m%d')}.docx"
    
    # 创建Word文档
    doc = Document()
    
    # 加载项目信息
    config = load_project_config(project_name)
    outline = load_outline(project_name)
    chapters = list_generated_chapters(project_name)
    
    # 添加标题
    doc.add_heading(config.get('name', project_name), level=0)
    
    # 添加大纲
    doc.add_heading('大纲', level=1)
    doc.add_paragraph(outline)
    
    # 添加正文
    doc.add_heading('正文', level=1)
    
    for chapter_file in chapters:
        chapter_content = load_chapter(project_name, chapter_file)
        doc.add_heading(chapter_file.replace('.md', ''), level=2)
        doc.add_paragraph(chapter_content)
    
    # 保存文档
    doc.save(output_path)
    
    return output_path


def export_to_epub(project_name, output_path=None):
    """导出为EPUB格式"""
    if output_path is None:
        output_path = f"{project_name}_{datetime.now().strftime('%Y%m%d')}.epub"
    
    # 创建EPUB书籍
    book = epub.EpubBook()
    
    # 加载项目信息
    config = load_project_config(project_name)
    outline = load_outline(project_name)
    chapters = list_generated_chapters(project_name)
    
    # 设置书籍元数据
    book.set_title(config.get('name', project_name))
    book.set_language('zh-CN')
    book.add_author('AI Novel Factory')
    
    # 创建章节
    epub_chapters = []
    
    # 添加大纲章节
    intro_chapter = epub.EpubHtml(title='大纲', file_name='intro.xhtml', lang='zh-CN')
    intro_chapter.content = f'<h1>大纲</h1><p>{outline.replace(chr(10), "<br/>")}</p>'
    book.add_item(intro_chapter)
    epub_chapters.append(intro_chapter)
    
    # 添加正文章节
    for i, chapter_file in enumerate(chapters):
        chapter_content = load_chapter(project_name, chapter_file)
        chapter_title = chapter_file.replace('.md', '')
        chapter = epub.EpubHtml(title=chapter_title, file_name=f'chapter_{i+1}.xhtml', lang='zh-CN')
        chapter.content = f'<h1>{chapter_title}</h1><p>{chapter_content.replace(chr(10), "<br/>")}</p>'
        book.add_item(chapter)
        epub_chapters.append(chapter)
    
    # 创建目录
    book.toc = epub_chapters
    
    # 添加默认的导航文件
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    
    # 设置书脊
    book.spine = ['nav'] + epub_chapters
    
    # 保存EPUB
    epub.write_epub(output_path, book, {})
    
    return output_path


def export_all_formats(project_name, output_dir=None):
    """导出所有格式"""
    if output_dir is None:
        output_dir = os.path.join('exports', project_name)
    
    # 创建输出目录
    os.makedirs(output_dir, exist_ok=True)
    
    # 导出各种格式
    txt_path = export_to_txt(project_name, os.path.join(output_dir, f"{project_name}.txt"))
    word_path = export_to_word(project_name, os.path.join(output_dir, f"{project_name}.docx"))
    epub_path = export_to_epub(project_name, os.path.join(output_dir, f"{project_name}.epub"))
    
    return {
        'txt': txt_path,
        'word': word_path,
        'epub': epub_path
    }
